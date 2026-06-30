"""Text extraction from .txt, .docx, and .pdf files.

Each extractor receives a BytesIO-like object and returns a plain Unicode string.
Extraction failures raise ValueError with a user-friendly message rather than
letting library exceptions bubble up raw.
"""
from __future__ import annotations

import io


def extract_text(data: io.BytesIO, ext: str) -> str:
    """Dispatch to the correct extractor based on file extension."""
    ext = ext.lower().lstrip(".")
    if ext == "txt":
        return _extract_txt(data)
    if ext == "docx":
        return _extract_docx(data)
    if ext == "pdf":
        return _extract_pdf(data)
    raise ValueError(f"Unsupported extension: .{ext}")


# ---------------------------------------------------------------------------
# .txt — plain read, try UTF-8 then fall back to Windows-1252
# ---------------------------------------------------------------------------

def _extract_txt(data: io.BytesIO) -> str:
    raw = data.read()
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    raise ValueError("Could not decode the text file. Make sure it is UTF-8 or a standard Windows encoding.")


# ---------------------------------------------------------------------------
# .docx — python-docx, join paragraph text with newlines
# ---------------------------------------------------------------------------

def _extract_docx(data: io.BytesIO) -> str:
    try:
        from docx import Document  # python-docx
    except ImportError:
        raise ValueError("python-docx is not installed. Run: pip install python-docx")

    try:
        doc = Document(data)
    except Exception as exc:
        raise ValueError(f"Could not open the Word document: {exc}")

    paragraphs = [p.text for p in doc.paragraphs]
    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.append(cell.text)

    text = "\n".join(p for p in paragraphs if p.strip())
    if not text.strip():
        raise ValueError("The Word document appears to be empty or contains no readable text.")
    return text


# ---------------------------------------------------------------------------
# .pdf — pdfplumber, concatenate page text
# ---------------------------------------------------------------------------

def _extract_pdf(data: io.BytesIO) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise ValueError("pdfplumber is not installed. Run: pip install pdfplumber")

    try:
        pages = []
        with pdfplumber.open(data) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    pages.append(page_text)
    except Exception as exc:
        raise ValueError(f"Could not read the PDF: {exc}")

    text = "\n\n".join(pages)
    if not text.strip():
        raise ValueError("The PDF appears to be empty, image-only, or scanned without OCR text.")
    return text
